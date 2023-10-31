import pyaudio
import wave
import keyboard  # Keyboard library is used to detect keypress
import speech_recognition as sr
import time
import os
import pyttsx3
import openai
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from os.path import expanduser
import sys


def record_audio(file_path):
    audio = pyaudio.PyAudio()

    format = pyaudio.paInt16
    channels = 1
    rate = 16000
    chunk = 1024
    input_device_index = None

    stream = audio.open(format=format, channels=channels, rate=rate,
                        input=True, input_device_index=input_device_index,
                        frames_per_buffer=chunk)

    print("Press SPACE to stop recording...")

    frames = []
    recording = True
    while recording:
        try:
            if keyboard.is_pressed('space'):  # Detect any key press
                recording = False
        except keyboard.KeyboardInterruptEvent:
            pass
        data = stream.read(chunk)
        frames.append(data)

    print("Finished recording.")

    stream.stop_stream()
    stream.close()
    audio.terminate()

    with wave.open(file_path, 'wb') as wf:
        wf.setnchannels(channels)
        wf.setsampwidth(audio.get_sample_size(format))
        wf.setframerate(rate)
        wf.writeframes(b''.join(frames))
        
def transcribe_audio(audio_file):
    recognizer = sr.Recognizer()

    with sr.AudioFile(audio_file) as source:
        audio_data = recognizer.record(source)  # Record the audio

    try:
        transcript = recognizer.recognize_google(audio_data)
        return transcript
    except sr.UnknownValueError:
        return " "
    except sr.RequestError as e:
        return f"Unfortunately the transcription engine has failed with the following error; {e}"
    
def question_and_response(question, audio_file_path):
    engine = pyttsx3.init()
    print(question)
    engine.say(question)
    print('Press SPACE to start recording')
    engine.say('Press SPACE to start recording, and to finish recording, please press SPACE again')
    engine.runAndWait()
    keyboard.read_key()
    keyboard.wait('space')
    record_audio(audio_file_path)
    transcript = transcribe_audio(audio_file_path)
    print("You said: " + transcript + ". Is this correct? If 'Yes', press 'Y'. If 'No' press SPACE to record again")
    engine.say("You said: " + transcript + ". Is this correct? If 'Yes', press 'Y'. If 'No' press SPACE to record again")
    engine.runAndWait()
    keyboard.read_key()
    if keyboard.is_pressed('Y'):
        correct = True
    elif keyboard.is_pressed('space'):
        correct = False
        while correct != True:
            engine.runAndWait()
            record_audio(audio_file_path)
            transcript = transcribe_audio(audio_file_path)
            print("You said: " + transcript + ". Is this correct? If 'Yes', press 'Y'. If 'No' press SPACE to record again")
            engine.say("You said: " + transcript + ". Is this correct? If 'Yes', press 'Y'. If 'No' press SPACE to record again" )
            engine.runAndWait()
            keyboard.read_key()
            if keyboard.is_pressed('Y'):
                correct = True
    return transcript

def generate_cover_letter(cv, job_description):
    print("I am now generating your cover letter, this may take a few moments")
    engine.say("I am now generating your cover letter, this may take a few moments")
    engine.runAndWait()
    # Set your OpenAI API key
    api_key = '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'
    
    openai.api_key = api_key
    
    cover_letter_conv = [
        {"role": "system", "content": "You are a helpful assistant that is great at writing cover letters for job applications."},
        {"role": "user", "content": f"Please write a cover letter for me for the position described in a job description provided based on the information that I have detailed in my CV."},
        {"role": "assistant", "content": "Sure, please provide your CV and the job description, and I'll help you write a cover letter to submit with your job application."},
        {"role": "user", "content": f"Here's my CV: {cv}"},
        {"role": "assistant", "content": "Great, now please provide the job description."},
        {"role": "user", "content": f"Here's the job description: {job_description}"}
    ]
    
    # Make the API request
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=cover_letter_conv,
    )
    
    # Extract the assistant's reply
    cover_letter = response['choices'][0]['message']['content']
    
    print(cover_letter)
    engine.say(cover_letter)
    engine.runAndWait()
    
    print("Are you happy with how this cover letter looks? If you would like to make some ammendments, press SPACE, otherwise, press 'Y'.")
    engine.say("Are you happy with how this cover letter looks? If you would like to make some ammendments, press SPACE, otherwise, press 'Y'.")
    engine.runAndWait()
    keyboard.read_key()
    if keyboard.is_pressed('Y'):
        cover_letter_correct = True
    elif keyboard.is_pressed('space'):
        cover_letter_correct = False
    while cover_letter_correct != True:
        cover_letter, cover_letter_conv = ammend_cover_letter(cover_letter_conv, cover_letter, cv, job_description)
        print(cover_letter)
        engine.say(cover_letter)
        engine.runAndWait()
        print("Are you happy with how this cover letter looks? If you would like to make some ammendments, press SPACE, otherwise, press 'Y'.")
        engine.say("Are you happy with how this cover letter looks? If you would like to make some ammendments, press SPACE, otherwise, press 'Y'.")
        engine.runAndWait()
        keyboard.read_key()
        if keyboard.is_pressed('Y'):
            cover_letter_correct = True
        elif keyboard.is_pressed('N'):
            cover_letter_correct = False
    cover_letter_to_docx(cover_letter, cv, job_description)
    return cover_letter

def compare_cv_job_description(cv, job_description):
    api_key = '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'
    
    openai.api_key = api_key
    print("I am just going to compare your CV to the job description and see how well matched you are for the position. This might take a moment")
    engine.say("I am just going to compare your CV to the job description and see how well matched you are for the position. This might take a moment")
    engine.runAndWait()
    
    conversation = [
        {"role": "system", "content": "You are a helpful assistant that provides feedback on a job application."},
        {"role": "user", "content": "I have a CV and a job description, and I need a summary of how well my skills mentioned in my CV match the job requirements detailed in the job description. Can you break these down into 'Strengths' and 'Shortcomings'."},
        {"role": "assistant", "content": "Sure, please provide the CV and job description, and I'll help you with the summary."},
        {"role": "user", "content": f"Here's my CV: {cv}"},
        {"role": "assistant", "content": "Great, now please provide the job description."},
        {"role": "user", "content": f"Here's the job description: {job_description}"}
    ]

    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=conversation
    )

    # Extract the AI's reply
    ai_reply = response['choices'][0]['message']['content']

    # Print the AI's response
    return ai_reply, conversation

def comparison_question(assistant_reply, messages, cv, job_description):
    api_key = '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'
    
    openai.api_key = api_key
    comparison_q = question_and_response("What would you like to know?", "recorded_audio.wav")
    messages += [
        {"role": "assistant", "content": assistant_reply},
        {"role": "user", "content": comparison_q}
    ]
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=messages,
    )
    assistant_reply = response['choices'][0]['message']['content']
    return assistant_reply, messages

def ammend_cover_letter(messages, assistant_reply, cv, job_description):
    api_key = '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'
    
    openai.api_key = api_key
    requested_ammendment = question_and_response("What changes would you like to make to the Cover Letter generated?", "recorded_audio.wav")
    messages += [
        {"role": "assistant", "content": assistant_reply},
        {"role": "user", "content": f"I need the following change(s) to be made to the cover letter: {requested_ammendment}"}
    ]
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=messages,
    )
    assistant_reply = response['choices'][0]['message']['content']
    return assistant_reply, messages

def get_company_name(job_description):
    api_key = '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'
    
    openai.api_key = api_key
    
    conversation = [
        {"role": "system", "content": "You are a helpful assistant that can read and return the name of the company that a job description is for."},
        {"role": "user", "content": "If I give you a position description, can you please extract the name of the company from the position description, and only return the name of the company in your response? Can you please answer the question by only returning the company name and no other text."},
        {"role": "assistant", "content": "Sure, please provide the job description, and I'll extract the company name for you."},
        {"role": "user", "content": f"Here's the job description: {job_description}"}
    ]

    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=conversation
    )

    # Extract the AI's reply
    ai_reply = response['choices'][0]['message']['content']

    # Print the AI's response
    return ai_reply

def get_applicant_name(cv):
    api_key = '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'
    
    openai.api_key = api_key
    
    conversation = [
        {"role": "system", "content": "You are a helpful assistant that can read and return the name of the person that a CV is for."},
        {"role": "user", "content": "If I give you a CV, can you please extract the name of the person that the CV is written about, and only return the name of the person in your response? Can you please answer the question by only returning the person's name and no other text."},
        {"role": "assistant", "content": "Sure, please provide the CV, and I'll extract the person's name for you."},
        {"role": "user", "content": f"Here's the job description: {cv}"}
    ]

    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=conversation
    )

    # Extract the AI's reply
    ai_reply = response['choices'][0]['message']['content']

    # Print the AI's response
    return ai_reply

def cover_letter_to_docx(cover_letter, cv, job_description):
    company = get_company_name(job_description)
    name = get_applicant_name(cv)
    print('Thank you. I will now save this Cover Letter to a word document in your Downloads folder')
    engine.say("Thank you. I will now save this Cover Letter to a word document in your Downloads folder")
    # Create a Word document and add the CV content
    doc = docx.Document()
    doc.add_heading("CV - " + name, 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(cv, style='Normal').alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Save the Word document to the Downloads folder
    downloads_folder = expanduser("~") + "/Downloads/"
    file_path = downloads_folder + "Cover Letter - " + name + company + ".docx"
    doc.save(file_path)
    print(f"Cover Letter has been saved to: {file_path}")
    engine.say(f"Cover Letter has been saved to: {file_path}")
    
    return cover_letter

def extract_docx(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

if __name__ == "__main__":
#def main():
    audio_file_path = "recorded_audio.wav"
    engine = pyttsx3.init()
    cover_letter_welcome_text = "Hello, welcome to the Cover Letter generator. In this section, I will ask you to upload a copy of your CV, as well as a copy of the job description for the role that you are interested in. In both cases, please enter the file path of each of your CV and the Job Description file into the prompts provided. Please also make sure that these files are not open on your computer when you enter the filepath. As well as generating a cover letter for you, I will do an assessment of how well your skills and experience match with that of the job you are interested in. From there, if there are any gaps in your CV or coverletter, we can ammend both these documents, or if there are skills or qualifications required that you need to attain, you can ask me for advice on how to attain them."
    print(cover_letter_welcome_text)
    engine.say(cover_letter_welcome_text)
    engine.runAndWait()
    engine.say("Please paste the path to the CV file in the space provided.")
    engine.runAndWait()
    cv_filepath = input('Please paste the path to the CV file:')
    #keyboard.read_key()
    cv = extract_docx(cv_filepath)
    engine.say("Please paste the path to the job description file in the space provided.")
    engine.runAndWait()
    job_description_filepath = input('Please paste the path to the job description file:')
    job_description = extract_docx(job_description_filepath)
    #keyboard.read_key()
    compare_output, compare_conv = compare_cv_job_description(cv, job_description)
    print(compare_output)
    engine.say(compare_output)
    engine.runAndWait()
    cover_letter = generate_cover_letter(cv, job_description)
    print("If you would like to ask questions or discuss the comparison that I gave, please press 1. To make updates to your CV, please press 2. If you would like to return to the home page, please press 3.")
    engine.say("If you would like to ask questions or discuss the comparison that I gave, please press 1. To make updates to your CV, please press 2. If you would like to return to the home page, please press 3.")
    engine.runAndWait()
    keyboard.read_key()
    if keyboard.is_pressed('3'):
        completed = True
    elif keyboard.is_pressed('2'):
        action = 'CV'
        completed = False
    elif keyboard.is_pressed('1'):
        action = 'Q'
        completed = False 
    while completed == False:
        if action == 'Q':
            print(compare_output)
            print(compare_conv)
            compare_output, compare_conv = comparison_question(compare_output, compare_conv, cv, job_description)
            print(compare_output)
            engine.say(compare_output)
            engine.runAndWait()
            print("If you would like to ask another question, press 'Y', otherwise press 'N'.")
            engine.say("If you would like to ask another question, press 'Y', otherwise press 'N'.")
            engine.runAndWait()
            keyboard.read_key()
            if keyboard.is_pressed('N'):
                completed = True
            elif keyboard.is_pressed('Y'):
                completed = False
        elif action == 'CV':
            cv = ammend_CV(cv)
            print(cv)
            print('Please see the generated CV above. If you need me to read this out to you, please press SPACE, otherwise, press any other key.')
            engine.say('Please see the generated CV above. If you need me to read this out to you, please press SPACE, otherwise, press any other key.')
            engine.runAndWait()
            keyboard.read_key()
            if keyboard.is_pressed('space'):
                engine.say('This is what the CV says: ' + cv)
            print("If you would like to make any further ammendments, press 'Y', otherwise press 'N'.")
            engine.say("If you would like to make any further ammendments, press 'Y', otherwise press 'N'.")
            engine.runAndWait()
            keyboard.read_key()
            if keyboard.is_pressed('N'):
                completed = True
            elif keyboard.is_pressed('Y'):
                completed = False
            
        