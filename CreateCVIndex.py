import pyaudio
import wave
import keyboard  # Keyboard library is used to detect keypress
import speech_recognition as sr
import time
import os
import pyttsx3
import openai
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from os.path import expanduser


def record_audio(file_path):
    audio = pyaudio.PyAudio()

    format = pyaudio.paInt16
    channels = 1
    rate = 16000
    chunk = 1024
    input_device_index = None

    stream = audio.open(format=format, channels=channels, rate=rate,
                        input=True, input_device_index=input_device_index,
                        frames_per_buffer=chunk)

    print("Press SPACE to stop recording...")

    frames = []
    recording = True
    while recording:
        try:
            if keyboard.is_pressed('space'):  # Detect any key press
                recording = False
        except keyboard.KeyboardInterruptEvent:
            pass
        data = stream.read(chunk)
        frames.append(data)

    print("Finished recording.")

    stream.stop_stream()
    stream.close()
    audio.terminate()

    with wave.open(file_path, 'wb') as wf:
        wf.setnchannels(channels)
        wf.setsampwidth(audio.get_sample_size(format))
        wf.setframerate(rate)
        wf.writeframes(b''.join(frames))
        
def transcribe_audio(audio_file):
    recognizer = sr.Recognizer()

    with sr.AudioFile(audio_file) as source:
        audio_data = recognizer.record(source)  # Record the audio

    try:
        transcript = recognizer.recognize_google(audio_data)
        return transcript
    except sr.UnknownValueError:
        return " "
    except sr.RequestError as e:
        return f"Unfortunately the transcription engine has failed with the following error; {e}"
    
def question_and_response(question, audio_file_path):
    engine = pyttsx3.init()
    print(question)
    engine.say(question)
    print('Press SPACE to start recording')
    engine.say('Press SPACE to start recording, and to finish recording, please press SPACE again')
    engine.runAndWait()
    keyboard.read_key()
    keyboard.wait('space')
    record_audio(audio_file_path)
    transcript = transcribe_audio(audio_file_path)
    print("You said: " + transcript + ". Is this correct? If 'Yes', press 'Y'. If 'No' press SPACE to record again")
    engine.say("You said: " + transcript + ". Is this correct? If 'Yes', press 'Y'. If 'No' press SPACE to record again")
    engine.runAndWait()
    keyboard.read_key()
    if keyboard.is_pressed('Y'):
        correct = True
    elif keyboard.is_pressed('space'):
        correct = False
        while correct != True:
            engine.runAndWait()
            record_audio(audio_file_path)
            transcript = transcribe_audio(audio_file_path)
            print("You said: " + transcript + ". Is this correct? If 'Yes', press 'Y'. If 'No' press SPACE to record again")
            engine.say("You said: " + transcript + ". Is this correct? If 'Yes', press 'Y'. If 'No' press SPACE to record again" )
            engine.runAndWait()
            keyboard.read_key()
            if keyboard.is_pressed('Y'):
                correct = True
    return transcript
    
def CV_generator(name, dob, address, phone, email, qualifications, work_experience, skills, interests, disabilities):
    engine = pyttsx3.init()
    print('Please wait while I generate you CV. This may take a moment.')
    engine.say('Please wait while I generate you CV. This may take a moment.')
    # Set your API key
    api_key = "INSERT YOUR OPEN AI API KEY HERE"
    openai.api_key = api_key

    # Define the conversation
    conversation = [
        {"role": "system", "content": "You are a helpful assistant that generates CVs."},
        {"role": "user", "content": "Please generate a CV based on the following information:"},
        {"role": "assistant", "content": """
         Name: """ + name + """
         DOB: """ + dob + """
         Address: """ + address + """
         Email: """ + email + """
         Phone: """ + phone + """
         Experience: """ + work_experience + """
         Skills: """ + skills + """
         Interests: """ + interests + """         
         Qualifications: """ + qualifications + """
         Disabilities: """ + disabilities + """
    """}
    ]

    # Generate the CV
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=conversation,
        )

    # Extract the generated CV
    cv = response['choices'][0]['message']['content']

    # Print the CV
    print(cv)
    print('Please see the generated CV above. If you need me to read this out to you, please press SPACE.')
    engine.say('Please see the generated CV above. If you need me to read this out to you, please press SPACE, otherwise, press any other key.')
    engine.runAndWait()
    keyboard.read_key()
    if keyboard.is_pressed('space'):
        engine.say('This is what the CV says: ' + cv)
    print("Does this look correct to you? If 'Yes', press 'Y'. If 'No' press 'N'.")
    engine.say("Does this look correct to you? If 'Yes', press 'Y'. If 'No' press 'N'.")
    engine.runAndWait()
    keyboard.read_key()
    if keyboard.is_pressed('Y'):
        cv_correct = True
    elif keyboard.is_pressed('N'):
        cv_correct = False
        while cv_correct != True:
            # Modify the CV with a location change
            transcript = question_and_response("What ammendment would you like to see to the CV generated?", "recorded_audio.wav")
            conversation += [{"role": "user", "content": transcript}]

            # Generate the modified CV
            modified_response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=conversation,
            )

            # Extract the modified CV
            cv = modified_response['choices'][0]['message']['content']

            # Print the modified CV
            print(cv)
            print('Please see the generated CV above. If you need me to read this out to you, please press SPACE.')
            engine.say('Please see the generated CV above. If you need me to read this out to you, please press SPACE, otherwise, press any other key.')
            engine.runAndWait()
            keyboard.read_key()
            if keyboard.is_pressed('space'):
                engine.say('This is what the CV says: ' + cv)
            print("Does this look correct to you? If 'Yes', press 'Y'. If 'No' press 'N'.")
            engine.say("Does this look correct to you? If 'Yes', press 'Y'. If 'No' press 'N'.")
            engine.runAndWait()
            keyboard.read_key()
            if keyboard.is_pressed('Y'):
                cv_correct = True
            elif keyboard.is_pressed('N'):
                cv_correct = False
    
    print('Thank you. I will now save this CV to a word document in your Downloads folder')
    engine.say("Thank you. I will now save this CV to a word document in your Downloads folder")
    # Create a Word document and add the CV content
    doc = docx.Document()
    doc.add_heading("CV - " + name, 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(cv, style='Normal').alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Save the Word document to the Downloads folder
    downloads_folder = expanduser("~") + "/Downloads/"
    file_path = downloads_folder + "CV - " + name + ".docx"
    doc.save(file_path)
    print(f"CV has been saved to: {file_path}")
    engine.say(f"CV has been saved to: {file_path}")
    
    return cv
    
#def multi_choice():

# NEED TO CREATE SEPARATE FUNCTION TO AMMEND CV

if __name__ == "__main__":
#def main():
    audio_file_path = "recorded_audio.wav"
    engine = pyttsx3.init()
    CV_welcome_text = "Hello, welcome to the CV generator. I am going to ask you a series of questions so that I can generate a CV for you."
    print(CV_welcome_text)
    engine.say(CV_welcome_text)
    engine.runAndWait()
    name = question_and_response("What is your full name?", audio_file_path)
    dob = question_and_response("What is your date of birth?", audio_file_path)
    address = question_and_response("What is your full address?", audio_file_path)
    phone = question_and_response("What is your best contact phone number?", audio_file_path)
    email = question_and_response("What is your email address?", audio_file_path)
    qualifications = question_and_response("Please give me a full list of your the qualifications that you have acheived. For each qualification, please also state the month and year of course commencement, the month and year of course completion, and the name of the institution you studied at. If you are still currently studying for a qualification, please say 'until Present'", audio_file_path)
    work_experience = question_and_response("Please give me details of your employment history. For each job, please also state the month and year you started working there, the month and year you finished working there, and the name of the organisation you worked for. If you are still currently employed in one of these positions, please say 'until Present'. If you have never been employed, please say 'No employment history'.", audio_file_path)
    skills = question_and_response("Please tell me about any other skills that you might have that would be relevant to your employment. These might include computer programs that you know how to use, areas of expertise or other soft skills that you might have developed.", audio_file_path)
    interests = question_and_response("Please tell me about your interests. These might include hobbies that you have or extra curricular activities that you participate in.", audio_file_path)
    disabilities = question_and_response("Please tell me about any disabilities that you might have. If you do not wish to disclose this information, please say 'Not Disclosed'.", audio_file_path)
    cv = CV_generator(name, dob, address, phone, email, qualifications, work_experience, skills, interests, disabilities)
    print("Your CV is now generated and ready to use! If you have a position description on-hand, we can now work together to build a cover letter or practice interviewing for this role. To do this, press 'Y'. Otherwise, press SPACE to return to the Home menu.")
    engine.say("Your CV is now generated and ready to use! If you have a position description on-hand, we can now work together to build a cover letter or practice interviewing for this role. To do this, press 'Y'. Otherwise, press any other key to return to the Home menu.")
    engine.runAndWait()
    keyboard.read_key()
    if keyboard.is_pressed('Y'):
        CreateCoverLetterIndex(cv)  
    
